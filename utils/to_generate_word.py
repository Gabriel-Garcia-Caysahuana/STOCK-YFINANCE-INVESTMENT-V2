# utils/to_generate_word.py

"""
Módulo para la generación de un informe en formato Word con tablas y gráficos
para el análisis de inversión.

Incluye:
  - Estadísticas descriptivas de los datos.
  - Gráficos de líneas, box plot de los datos.
  - Pesos óptimos de cada ticker en el portafolio.
  - NUEVO: Secciones adicionales para correlación, volatilidad móvil e histogramas de retornos.
"""

import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from typing import Optional, Dict, List
import pandas as pd

from .analyse_data import (
    descriptive, 
    plot_line_series, 
    plot_box_plot,
    get_correlation_matrix,
    plot_correlation_heatmap,
    plot_rolling_volatility,
    plot_returns_histogram
)

def generar_informe(
    data: pd.DataFrame,
    tickers: List[str],
    weights: Dict[str, float],
    output: Optional[str] = None,
    window_volatility: int = 20
) -> None:
    """
    Genera un informe de análisis de inversión en formato Word, incluyendo:
      - Estadísticas descriptivas de los datos.
      - Gráficos de líneas y box plot de los datos.
      - Pesos óptimos de cada ticker en el portafolio.
      - Matriz de correlación y heatmap.
      - Volatilidad móvil (rolling) para cada ticker.
      - Histogramas de retornos para cada ticker.

    :param data: DataFrame con precios (y retornos, si existen) de los activos.
    :param tickers: Lista de símbolos de acciones a analizar.
    :param weights: Diccionario con los pesos óptimos de cada ticker.
    :param output: Ruta completa del archivo de salida. Si no se especifica,
                   se guarda un archivo 'informe_analisis_inversion.docx' 
                   en el directorio actual.
    :param window_volatility: Ventana (en días) para calcular la volatilidad móvil
                              de cada ticker (por defecto 20).
    :return: None
    """
    if data.empty:
        print("El DataFrame proporcionado está vacío. No se puede generar el informe.")
        return

    doc = Document()
    doc.add_heading('Informe de Análisis de Inversión', level=1)
    doc.add_paragraph(
        "Este informe contiene un análisis de los pesos óptimos del portafolio, "
        "estadísticas descriptivas de los activos y diversos gráficos visuales "
        "de su rendimiento y riesgo."
    )

    # -------------------------------------------------------------------------
    # 1. Estadísticas descriptivas
    # -------------------------------------------------------------------------
    doc.add_heading('Estadísticas Descriptivas', level=2)
    stats = descriptive(data)

    if not stats.empty:
        _agregar_tabla_estadisticas(doc, stats)
    else:
        doc.add_paragraph("No se pudo generar estadísticas descriptivas porque el DataFrame está vacío.")

    # -------------------------------------------------------------------------
    # 2. Gráficos de rendimiento (Series temporales, Box Plot)
    # -------------------------------------------------------------------------
    doc.add_heading('Gráficos de Rendimiento', level=2)

    fig1_path = "temp_plot_line_series.png"
    if _generar_y_guardar_grafico(plot_line_series, data, tickers, fig1_path):
        doc.add_paragraph("Gráfico de Series Temporales:")
        doc.add_picture(fig1_path, width=Inches(6))

    fig2_path = "temp_plot_box_plot.png"
    if _generar_y_guardar_grafico(plot_box_plot, data, tickers, fig2_path):
        doc.add_paragraph("Gráfico de Box Plot:")
        doc.add_picture(fig2_path, width=Inches(6))

    # -------------------------------------------------------------------------
    # 3. Pesos óptimos del portafolio
    # -------------------------------------------------------------------------
    doc.add_heading('Pesos Óptimos del Portafolio', level=2)
    if weights:
        for ticker, weight in weights.items():
            doc.add_paragraph(f"{ticker}: {float(weight):.2%}")
    else:
        doc.add_paragraph("No se encontraron pesos óptimos para el portafolio.")

    # -------------------------------------------------------------------------
    # 4. Sección de Análisis Adicional
    # -------------------------------------------------------------------------
    doc.add_heading('Análisis Adicional', level=2)

    # 4.1. Correlación
    doc.add_heading('Correlación de Retornos', level=3)
    _crear_seccion_correlacion(doc, data, tickers)

    # 4.2. Volatilidad móvil (por ticker)
    doc.add_heading('Volatilidad Móvil (Rolling)', level=3)
    _crear_seccion_volatilidad(doc, data, tickers, window=window_volatility)

    # 4.3. Histogramas de Retornos
    doc.add_heading('Histogramas de Retornos', level=3)
    _crear_seccion_histogramas(doc, data, tickers)

    # -------------------------------------------------------------------------
    # Limpieza de archivos temporales
    # -------------------------------------------------------------------------
    _eliminar_archivo_temp(fig1_path)
    _eliminar_archivo_temp(fig2_path)

    # Guardar documento final
    if output:
        doc.save(output)
        file_path = os.path.abspath(output)
        print(f"El informe de análisis ha sido guardado como '{output}' en '{file_path}'")
    else:
        doc_path = 'informe_analisis_inversion.docx'
        doc.save(doc_path)
        file_path = os.path.abspath(doc_path)
        print(f"El informe de análisis ha sido guardado como '{doc_path}' en '{file_path}'")


# ============================================================================
# Funciones Auxiliares
# ============================================================================

def _agregar_tabla_estadisticas(doc: Document, stats: pd.DataFrame) -> None:
    """
    Función auxiliar para agregar una tabla de estadísticas descriptivas al documento.
    """
    # Crear tabla con número de columnas = número de columnas en stats + 1 para la fila de índice
    table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
    hdr_cells = table.rows[0].cells

    # Encabezados de la tabla
    hdr_cells[0].text = 'Ticker'
    for i, column in enumerate(stats.columns):
        hdr_cells[i + 1].text = str(column)

    # Filas de la tabla
    for row_idx, row_data in stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_idx)
        for i, value in enumerate(row_data):
            row_cells[i + 1].text = str(value)


def _generar_y_guardar_grafico(
    plot_function,
    data: pd.DataFrame,
    tickers: list,
    output_path: str
) -> bool:
    """
    Función auxiliar para generar y guardar un gráfico en formato PNG, utilizando
    una función de plotting específica (plot_function).

    :param plot_function: Función que crea el gráfico (p.ej. plot_line_series).
    :param data: DataFrame con la información a graficar.
    :param tickers: Lista de símbolos de acciones a graficar.
    :param output_path: Ruta del archivo PNG donde se guardará el gráfico.
    :return: True si el archivo se generó exitosamente, False en caso contrario.
    """
    # Borramos gráficos anteriores de la memoria
    plt.close('all')

    try:
        plot_function(data, tickers)
        plt.savefig(output_path, format='png', bbox_inches='tight')
        plt.close()
        return True
    except Exception as e:
        print(f"No se pudo generar el gráfico '{output_path}'. Error: {e}")
        plt.close()
        return False


def _eliminar_archivo_temp(file_path: str) -> None:
    """
    Función auxiliar para eliminar un archivo temporal, si existe.
    """
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Error al eliminar el archivo temporal '{file_path}': {e}")


# ============================================================================
# NUEVAS SECCIONES: Correlación, Volatilidad móvil, Histogramas
# ============================================================================
def _crear_seccion_correlacion(doc: Document, data: pd.DataFrame, tickers: List[str]) -> None:
    """
    Crea una sección en el documento con la matriz de correlación de los retornos
    y un heatmap de la misma.
    """
    # Preparamos las columnas de retornos
    ret_cols = [f'R{t}' for t in tickers if f'R{t}' in data.columns]
    if not ret_cols:
        doc.add_paragraph("No hay columnas de retornos para generar la matriz de correlación.")
        return

    corr_matrix = get_correlation_matrix(data, ret_cols)
    if corr_matrix.empty:
        doc.add_paragraph("No se pudo generar la matriz de correlación. No hay datos válidos.")
        return

    # Agregamos tabla con la matriz de correlación (opcional)
    doc.add_paragraph("Matriz de correlación (valores numéricos):")
    _agregar_tabla_correlacion(doc, corr_matrix)

    # Generamos heatmap y lo insertamos
    fig_corr_path = "temp_correlation_heatmap.png"
    plt.close('all')
    try:
        plot_correlation_heatmap(corr_matrix, title="Heatmap de Correlación de Retornos")
        plt.savefig(fig_corr_path, format='png', bbox_inches='tight')
        plt.close()
        doc.add_paragraph("Mapa de Calor (Heatmap):")
        doc.add_picture(fig_corr_path, width=Inches(5))
    except Exception as e:
        doc.add_paragraph(f"No se pudo generar el heatmap de correlación. Error: {e}")

    _eliminar_archivo_temp(fig_corr_path)


def _agregar_tabla_correlacion(doc: Document, corr_matrix: pd.DataFrame) -> None:
    """
    Agrega una tabla al documento con la matriz de correlación.
    """
    # Encabezados de filas y columnas
    tickers = corr_matrix.columns.tolist()
    table = doc.add_table(rows=1, cols=len(tickers) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = " "  # Celda vacía en la esquina superior izquierda

    # Llenamos la fila de cabecera con los nombres de columnas
    for i, col in enumerate(tickers):
        hdr_cells[i + 1].text = col

    # Llenamos las filas con datos
    for row_idx, row_name in enumerate(tickers):
        row_cells = table.add_row().cells
        row_cells[0].text = row_name  # Nombre del ticker en la primera columna
        for col_idx, col_name in enumerate(tickers):
            valor = corr_matrix.loc[row_name, col_name]
            row_cells[col_idx + 1].text = f"{valor:.2f}"


def _crear_seccion_volatilidad(
    doc: Document, 
    data: pd.DataFrame, 
    tickers: List[str], 
    window: int
) -> None:
    """
    Crea una sección en el documento mostrando la volatilidad móvil para cada ticker.
    """
    for ticker in tickers:
        fig_vol_path = f"temp_rolling_volatility_{ticker}.png"
        plt.close('all')
        try:
            plot_rolling_volatility(
                data, 
                ticker=ticker, 
                window=window, 
                title=f"Volatilidad Móvil ({window} días) - {ticker}"
            )
            plt.savefig(fig_vol_path, format='png', bbox_inches='tight')
            plt.close()
            doc.add_paragraph(f"Volatilidad móvil para {ticker}:")
            doc.add_picture(fig_vol_path, width=Inches(5))
        except Exception as e:
            doc.add_paragraph(f"No se pudo generar la volatilidad móvil para {ticker}. Error: {e}")
        _eliminar_archivo_temp(fig_vol_path)


def _crear_seccion_histogramas(
    doc: Document,
    data: pd.DataFrame,
    tickers: List[str]
) -> None:
    """
    Crea una sección en el documento con el histograma de retornos de cada ticker.
    """
    for ticker in tickers:
        fig_hist_path = f"temp_histogram_{ticker}.png"
        plt.close('all')
        try:
            plot_returns_histogram(
                data, 
                ticker=ticker, 
                bins=30, 
                title=f"Histograma de Retornos - {ticker}"
            )
            plt.savefig(fig_hist_path, format='png', bbox_inches='tight')
            plt.close()
            doc.add_paragraph(f"Histograma de retornos para {ticker}:")
            doc.add_picture(fig_hist_path, width=Inches(5))
        except Exception as e:
            doc.add_paragraph(f"No se pudo generar el histograma de retornos para {ticker}. Error: {e}")
        _eliminar_archivo_temp(fig_hist_path)
